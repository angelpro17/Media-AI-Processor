import re
import os
import logging
import httpx
from typing import Dict, Any, List, Tuple

log = logging.getLogger(__name__)

SUPPORTED_PAIRS = [
    "es-en",
    "en-es",
    "es-fr",
    "fr-es",
    "es-it",
    "it-es",
    "es-pt",
    "pt-es",
    "es-de",
    "de-es",
]

# MyMemory API — free up to 10k words/day, no API key needed
_MYMEMORY_URL = "https://api.mymemory.translated.net/get"

def _translate_via_api(text: str, src: str, tgt: str) -> str:
    """Translate a single chunk of text via MyMemory free API."""
    try:
        resp = httpx.get(
            _MYMEMORY_URL,
            params={"q": text, "langpair": f"{src}|{tgt}"},
            timeout=15.0,
        )
        data = resp.json()
        translated = data.get("responseData", {}).get("translatedText", "")
        if translated and data.get("responseStatus") == 200:
            return translated
        # Fallback: try matches list
        matches = data.get("matches", [])
        if matches:
            return matches[0].get("translation", text)
        return text
    except Exception as e:
        log.warning("MyMemory API error: %s — returning original text", e)
        return text

# For ROMANCE models, MarianMT requires a target language prefix like ">>pt<<"
prefix_map = {
    "pt": ">>pt<<",
    "fr": ">>fr<<",
    "es": ">>es<<",
    "it": ">>it<<"
}

_cache: Dict[str, Any] = {}

MAX_BATCH_CHARS = 400
MAX_BATCH_SIZE  = 32

def _get_pipeline(direction: str) -> Dict[str, Any]:
    from transformers import MarianMTModel, MarianTokenizer
    model_name = MODEL_MAP.get(direction)
    if not model_name:
        raise ValueError(f"No translation model defined for '{direction}'")
        
    if model_name not in _cache:
        log.info("Loading translation model: %s", model_name)
        
        import torch
        device = "cpu"
        if torch.backends.mps.is_available():
            device = "mps"
        elif torch.cuda.is_available():
            device = "cuda"
            
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model     = MarianMTModel.from_pretrained(model_name).to(device)
        model.eval()
        _cache[model_name] = {"model": model, "tokenizer": tokenizer, "device": device}
        log.info("Model %s loaded on device: %s", model_name, device)
    return _cache[model_name]

def _batch_translate(texts: List[str], direction: str) -> List[str]:
    """Translate a batch of texts — uses API when on Render, falls back to fast local model."""
    if not texts:
        return []

    src_lang, tgt_lang = direction.split("-")
    
    # If running on Render (Low RAM/CPU server), default to API to prevent heavy CPU loads
    if os.getenv("RENDER"):
        return [_translate_via_api(t, src_lang, tgt_lang) for t in texts]

    # Offline/Local GPU mode
    model_name = MODEL_MAP[direction]
    pipeline  = _get_pipeline(direction)
    tokenizer = pipeline["tokenizer"]
    model     = pipeline["model"]
    device = pipeline["device"]

    # Prepend target language prefix if using a multilingual ROMANCE model
    if "ROMANCE" in model_name:
        prefix = prefix_map.get(tgt_lang, f">>{tgt_lang}<<")
        processed_texts = [f"{prefix} {t}" for t in texts]
    else:
        processed_texts = texts

    results: List[str] = []
    for i in range(0, len(processed_texts), MAX_BATCH_SIZE):
        batch  = processed_texts[i : i + MAX_BATCH_SIZE]
        inputs = tokenizer(batch, return_tensors="pt", padding=True, truncation=True, max_length=512)
        inputs = {k: v.to(device) for k, v in inputs.items()}
        outputs = model.generate(**inputs, num_beams=4, max_length=512)
        decoded = tokenizer.batch_decode(outputs, skip_special_tokens=True)
        results.extend(decoded)
    return results


def _split_into_chunks(text: str, max_chars: int = 400) -> List[str]:
    """Split a paragraph into sentence chunks."""
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    chunks: List[str] = []
    current = ""
    for sent in sentences:
        if len(current) + len(sent) > max_chars and current:
            chunks.append(current.strip())
            current = sent
        else:
            current = (current + " " + sent).strip()
    if current:
        chunks.append(current)
    return chunks or [text]


def translate_text(text: str, direction: str) -> str:
    """Translate plain text using MyMemory API, preserving paragraph structure."""
    src, tgt = direction.split("-")
    paragraphs = text.split("\n")
    out_paragraphs: List[str] = []

    for para in paragraphs:
        stripped = para.strip()
        if not stripped:
            out_paragraphs.append("")
            continue
        chunks = _split_into_chunks(stripped)
        translated_chunks = [_translate_via_api(c, src, tgt) for c in chunks]
        out_paragraphs.append(" ".join(translated_chunks))

    return "\n".join(out_paragraphs)


# ── Legacy local-model helpers (only used for document translation) ──────────
MODEL_MAP = {
    "es-fr": "Helsinki-NLP/opus-mt-es-ROMANCE",
    "es-it": "Helsinki-NLP/opus-mt-es-ROMANCE",
    "es-pt": "Helsinki-NLP/opus-mt-es-ROMANCE",
    "en-es": "Helsinki-NLP/opus-mt-en-ROMANCE",
    "en-fr": "Helsinki-NLP/opus-mt-en-ROMANCE",
    "es-en": "Helsinki-NLP/opus-mt-es-en",
    "fr-es": "Helsinki-NLP/opus-mt-fr-es",
    "it-es": "Helsinki-NLP/opus-mt-it-es",
    "pt-es": "Helsinki-NLP/opus-mt-pt-es",
    "es-de": "Helsinki-NLP/opus-mt-es-de",
    "de-es": "Helsinki-NLP/opus-mt-de-es",
}
prefix_map = {"pt": ">>pt<<", "fr": ">>fr<<", "es": ">>es<<", "it": ">>it<<"}
_cache: Dict[str, Any] = {}
MAX_BATCH_CHARS = 400
MAX_BATCH_SIZE  = 32


def _get_pipeline(direction: str) -> Dict[str, Any]:
    from transformers import MarianMTModel, MarianTokenizer
    model_name = MODEL_MAP.get(direction)
    if not model_name:
        raise ValueError(f"No translation model defined for '{direction}'")
    if model_name not in _cache:
        log.info("Loading translation model: %s", model_name)
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model = MarianMTModel.from_pretrained(model_name)
        model.eval()
        _cache[model_name] = {"model": model, "tokenizer": tokenizer}
    return _cache[model_name]


def _batch_translate(texts: List[str], direction: str) -> List[str]:
    """Translate a batch of texts — uses API when possible, falls back to local model."""
    if not texts:
        return []
    src, tgt = direction.split("-")
    # Use API for short texts (document translation chunks)
    return [_translate_via_api(t, src, tgt) for t in texts]



# ── PDF → PDF (batched, in-place layout preservation) ─────────
def translate_pdf(src: str, direction: str, out_path: str) -> str:
    """
    Translate a PDF preserving its exact visual layout.
    All spans across each page are batched in a single model call.
    """
    import fitz

    doc = fitz.open(src)
    total_pages = len(doc)

    for page_idx, page in enumerate(doc):
        log.info("Translating page %d/%d", page_idx + 1, total_pages)
        page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        # Collect span metadata and raw texts
        span_meta: List[Dict] = []
        raw_texts:  List[str] = []

        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    span_meta.append({
                        "bbox":   fitz.Rect(span["bbox"]),
                        "origin": fitz.Point(span["origin"]),
                        "size":   span["size"],
                        "color":  span.get("color", 0),
                    })
                    raw_texts.append(text)

        if not raw_texts:
            continue

        # Single batch translate for the whole page
        translated_texts = _batch_translate(raw_texts, direction)

        # Step 1: redact originals
        for meta in span_meta:
            page.add_redact_annot(meta["bbox"], fill=(1, 1, 1), text="")
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        # Step 2: reinsert translations
        for meta, translated in zip(span_meta, translated_texts):
            c = meta["color"]
            color = (((c >> 16) & 0xFF) / 255, ((c >> 8) & 0xFF) / 255, (c & 0xFF) / 255)
            try:
                page.insert_text(meta["origin"], translated,
                                  fontsize=meta["size"], color=color, overlay=True)
            except Exception:
                page.insert_text(meta["origin"], translated, fontsize=meta["size"])

    doc.save(out_path, garbage=4, deflate=True)
    doc.close()
    return out_path


# ── DOCX → DOCX (batched paragraph translation) ───────────────
def translate_docx(src: str, direction: str, out_path: str) -> str:
    """
    Translate a DOCX preserving paragraph formatting.
    All paragraphs are batched and translated in a single pass.
    """
    from docx import Document

    doc = Document(src)

    # Collect all texts to translate (paragraphs + table cells)
    # Each entry: (text, setter_fn)
    items: List[Tuple[str, Any]] = []

    def _collect_para(para):
        text = para.text.strip()
        if not text:
            return

        def _setter(translated: str):
            if para.runs:
                para.runs[0].text = translated
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = translated

        items.append((text, _setter))

    for para in doc.paragraphs:
        _collect_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _collect_para(para)

    if items:
        texts   = [t for t, _ in items]
        setters = [s for _, s in items]
        translated_all = _batch_translate(texts, direction)
        for setter, translated in zip(setters, translated_all):
            setter(translated)

    doc.save(out_path)
    return out_path


# ── TXT → TXT ─────────────────────────────────────────────────
def translate_txt(src: str, direction: str, out_path: str) -> str:
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    translated = translate_text(text, direction)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(translated)
    return out_path


# ── SRT / VTT → SRT / VTT ─────────────────────────────────────
def translate_subs(src: str, direction: str, out_path: str) -> str:
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()

    # Split by double newline to get subtitle blocks
    blocks = content.strip().split("\n\n")
    
    texts_to_translate = []
    block_headers = []
    
    for block in blocks:
        lines = block.split("\n")
        
        # A typical SRT block has at least 3 lines: Index, Timestamp, Text.
        # A VTT might have WEBVTT at the start, or 2 lines: Timestamp, Text.
        header_lines = []
        text_lines = []
        
        for line in lines:
            if "-->" in line or line.strip().isdigit() or "WEBVTT" in line:
                header_lines.append(line)
            else:
                text_lines.append(line)
                
        texts_to_translate.append("\n".join(text_lines))
        block_headers.append(header_lines)

    translated_texts = _batch_translate(texts_to_translate, direction)
    
    out_blocks = []
    for header, translated in zip(block_headers, translated_texts):
        if header:
            if translated.strip():
                out_blocks.append("\n".join(header) + "\n" + translated)
            else:
                out_blocks.append("\n".join(header))
        else:
            out_blocks.append(translated)
            
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(out_blocks) + "\n")
        
    return out_path


# ── Background task dispatcher ────────────────────────────────
def run_document_translation(job_id: str, src_path: str, ext: str,
                              direction: str, out_name: str) -> None:
    import os, tempfile
    from core.jobs import update_job

    out_dir = tempfile.gettempdir()

    try:
        update_job(job_id, status="processing", progress=10)

        if ext in ("docx", "doc"):
            out_path = os.path.join(out_dir, f"{job_id}_translated.docx")
            translate_docx(src_path, direction, out_path)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        elif ext == "pdf":
            out_path = os.path.join(out_dir, f"{job_id}_translated.pdf")
            translate_pdf(src_path, direction, out_path)
            mime     = "application/pdf"
            out_name = out_name.rsplit(".", 1)[0] + "_traducido.pdf"

        elif ext == "txt":
            out_path = os.path.join(out_dir, f"{job_id}_translated.txt")
            translate_txt(src_path, direction, out_path)
            mime = "text/plain"

        elif ext in ("srt", "vtt"):
            out_path = os.path.join(out_dir, f"{job_id}_translated.{ext}")
            translate_subs(src_path, direction, out_path)
            mime = "text/plain"

        else:
            raise ValueError(f"Formato no soportado para traducción: {ext}")

        update_job(job_id, status="done", progress=100,
                   result_path=out_path, filename=out_name, mime_type=mime)
        log.info("Doc translation job %s done → %s", job_id, out_path)

    except Exception as e:
        log.exception("Doc translation job %s failed", job_id)
        update_job(job_id, status="error", error=str(e))
    finally:
        try:
            os.remove(src_path)
        except OSError:
            pass
